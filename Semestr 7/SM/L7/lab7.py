import numpy as np
import matplotlib.pyplot as plt
import cv2
import scipy.fftpack
from docx import Document
from docx.shared import Inches
from io import BytesIO
import sys
import humanize

def get_size(obj, seen=None):
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    seen.add(obj_id)
    if isinstance(obj, np.ndarray):
        size = obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size

def get_human_size(size):
    return humanize.naturalsize(size)

class ver1:
  Y=np.array([])
  Cb=np.array([])
  Cr=np.array([])
  ChromaRatio="4:4:4"
  QY=np.ones((8,8))
  QC=np.ones((8,8))
  shape=(0,0,3)

QY = np.array([
        [16, 11, 10, 16, 24,  40,  51,  61],
        [12, 12, 14, 19, 26,  58,  60,  55],
        [14, 13, 16, 24, 40,  57,  69,  56],
        [14, 17, 22, 29, 51,  87,  80,  62],
        [18, 22, 37, 56, 68,  109, 103, 77],
        [24, 36, 55, 64, 81,  104, 113, 92],
        [49, 64, 78, 87, 103, 121, 120, 101],
        [72, 92, 95, 98, 112, 100, 103, 99],
        ])

QC = np.array([
        [17, 18, 24, 47, 99, 99, 99, 99],
        [18, 21, 26, 66, 99, 99, 99, 99],
        [24, 26, 56, 99, 99, 99, 99, 99],
        [47, 66, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        ])

QN = np.ones((8,8))



def rle_encode(img):
  img = img.astype(int)
  img_shape = img.shape
  img_flatten = img.flatten()
  sign_counter = 0
  previous_sign = None
  coded_stream = []
  for sign in img_flatten:
    if sign == previous_sign:
      sign_counter += 1
    else:
      if previous_sign is not None:
        coded_stream.append(sign_counter)
        coded_stream.append(previous_sign)
      sign_counter = 1
      previous_sign = sign
  coded_stream.append(sign_counter)
  coded_stream.append(previous_sign)
  coded_stream = np.array(coded_stream)

  return [img_shape, coded_stream]

def rle_decode(coded_data):
  coded_shape = coded_data[0]
  coded_stream = coded_data[1]
  decoded_stream = []
  for i in range(0, len(coded_stream), 2):
    sign = coded_stream[i + 1]
    sign_counter = coded_stream[i]
    for j in range(sign_counter):
      decoded_stream.append(sign)
  decoded_stream = np.array(decoded_stream)
  decoded_stream = decoded_stream.reshape(coded_shape)

  return decoded_stream

def zigzag(A):
  template = np.array([
          [0,  1,  5,  6,  14, 15, 27, 28],
          [2,  4,  7,  13, 16, 26, 29, 42],
          [3,  8,  12, 17, 25, 30, 41, 43],
          [9,  11, 18, 24, 31, 40, 44, 53],
          [10, 19, 23, 32, 39, 45, 52, 54],
          [20, 22, 33, 38, 46, 51, 55, 60],
          [21, 34, 37, 47, 50, 56, 59, 61],
          [35, 36, 48, 49, 57, 58, 62, 63],
          ])
  if len(A.shape)==1:
    B=np.zeros((8,8))
    for r in range(0,8):
      for c in range(0,8):
        B[r,c]=A[template[r,c]]
  else:
    B=np.zeros((64,))
    for r in range(0,8):
      for c in range(0,8):
        B[template[r,c]]=A[r,c]
  return B

def chromaSubsampling(data):
  # wariant 4:2:2 - czyli redukujemy co drugą kolumnę
  reduced = np.empty((data.shape[0], int(data.shape[1]/2)))
  for row in range(0, data.shape[0]):
    for col in range(0, data.shape[1], 2):
      reduced[row][int(col/2)] = data[row][col]
  return reduced

def chromaResampling(data):
  # wariant 4:2:2 - czyli redukcja ilości pikseli o 50% w poziomie
  # -> każda wartość będzie odtwarzana za pomocą piksela na lewo od niego
  recreated = np.empty((data.shape[0], data.shape[1]*2))
  for row in range(0, data.shape[0]):
    for col in range(0, data.shape[1]):
      recreated[row][col*2] = data[row][col]
      recreated[row][col*2+1] = data[row][col]
  return recreated

def quantization(data, q_table):
    quantized = np.zeros((data.shape[0], data.shape[1]))
    index = 0
    for row in range(0, data.shape[0], 8):
        for col in range(0, data.shape[1], 8):
            quantized[index:index+64] = np.round(data/q_table).astype(int)
            index += 64
    return quantized


def dequantization(data, q_table):
    for idx, i in enumerate(range(0, data.shape[0], 8)):
        dequantized = data[i:i + 8] * q_table
    return dequantized

def dataToBlocks_8x8(data_color):
  data_blocks_8x8 = []
  for row in range(0, data_color.shape[0], 8):
    for col in range(0, data_color.shape[1], 8):
      block = data_color[row:row + 8, col:col + 8]
      data_blocks_8x8.append(block)
  return data_blocks_8x8

def blocks_8x8_ToData(blocks, sizeA, sizeB):
    data = np.zeros((sizeA, sizeB))
    index = 0
    for row in range(0, sizeA, 8):
        for col in range(0, sizeB, 8):
            data[row:row + 8, col:col + 8] = blocks[index]
            index += 1
    return data

def dct2(a):
  return scipy.fftpack.dct( scipy.fftpack.dct( a.astype(float), axis=0, norm='ortho' ), axis=1, norm='ortho' )

def idct2(a):
  return scipy.fftpack.idct( scipy.fftpack.idct( a.astype(float), axis=0 , norm='ortho'), axis=1 , norm='ortho')

def CompressLayer(data, QX):
  data = dataToBlocks_8x8(data)

  for i in range(len(data)):
    data[i] = data[i].astype(int) - 128
    data[i] = dct2(data[i])
    data[i] = quantization(data[i], QX)
    data[i] = zigzag(data[i])

  return data

def DecompressLayer(data, QX, ratio):
  temp_data = []

  for i in range(data.shape[0]):
    temp_data.append(data[i])

  for i in range(len(temp_data)):
    temp_data[i] = zigzag(temp_data[i])
    temp_data[i] = dequantization(temp_data[i], QX)
    temp_data[i] = idct2(temp_data[i]) + 128

  if ratio == "4:2:2":
    temp_data = blocks_8x8_ToData(temp_data, int(np.sqrt(len(temp_data)*2) * 8), int(np.sqrt(len(temp_data)*2) * 8/2))
  else:
    temp_data = blocks_8x8_ToData(temp_data, int(np.sqrt(len(temp_data)) * 8), int(np.sqrt(len(temp_data)) * 8))

  return temp_data

def CompressJPEG(RGB,Ratio="4:4:4",QY=np.ones((8,8)),QC=np.ones((8,8))):
  # RGB -> YCrCb
  img_YCrCb = cv2.cvtColor(RGB, cv2.COLOR_RGB2YCrCb).astype(int)
  JPEG= ver1()
  JPEG.shape = (1,1)

  # zapisać dane z wejścia do klasy
  JPEG.Y, JPEG.Cr, JPEG.Cb = np.clip(cv2.split(img_YCrCb), 0, 255)
  JPEG.ChromaRatio = Ratio
  JPEG.QY = QY
  JPEG.QC = QC

  # Tu chroma subsampling
  if Ratio != "4:4:4":
    JPEG.Y = chromaSubsampling(JPEG.Y)
    JPEG.Cr = chromaSubsampling(JPEG.Cr)
    JPEG.Cb = chromaSubsampling(JPEG.Cb)

  JPEG.Y=CompressLayer(JPEG.Y,JPEG.QY)
  JPEG.Cr=CompressLayer(JPEG.Cr,JPEG.QC)
  JPEG.Cb=CompressLayer(JPEG.Cb,JPEG.QC)

  # tu dochodzi kompresja bezstratna
  Y_temp = np.zeros((len(JPEG.Y), len(JPEG.Y[0])))
  Cr_temp = np.zeros((len(JPEG.Cr), len(JPEG.Cr[0])))
  Cb_temp = np.zeros((len(JPEG.Cb), len(JPEG.Cb[0])))

  for i in range(len(JPEG.Y)):
    Y_temp[i, :] = JPEG.Y[i]
    Cr_temp[i, :] = JPEG.Cr[i]
    Cb_temp[i, :] = JPEG.Cb[i]

  JPEG.Y = rle_encode(Y_temp)
  JPEG.Cr = rle_encode(Cr_temp)
  JPEG.Cb = rle_encode(Cb_temp)

  return JPEG

def DecompressJPEG(JPEG):
  # dekompresja bezstratna
  JPEG.Y = rle_decode(JPEG.Y)
  JPEG.Cr = rle_decode(JPEG.Cr)
  JPEG.Cb = rle_decode(JPEG.Cb)

  JPEG.Y=DecompressLayer(JPEG.Y,JPEG.QY, JPEG.ChromaRatio)
  JPEG.Cr=DecompressLayer(JPEG.Cr,JPEG.QC, JPEG.ChromaRatio)
  JPEG.Cb=DecompressLayer(JPEG.Cb,JPEG.QC, JPEG.ChromaRatio)

  # Tu chroma resampling
  if JPEG.ChromaRatio != "4:4:4":
    JPEG.Y = chromaResampling(JPEG.Y)
    JPEG.Cr = chromaResampling(JPEG.Cr)
    JPEG.Cb = chromaResampling(JPEG.Cb)

  # tu rekonstrukcja obrazu
  JPEG.Y = np.clip(JPEG.Y, 0, 255).astype(np.uint8)
  JPEG.Cr = np.clip(JPEG.Cr, 0, 255).astype(np.uint8)
  JPEG.Cb = np.clip(JPEG.Cb, 0, 255).astype(np.uint8)

  YCrCb2 = np.dstack([JPEG.Y, JPEG.Cr, JPEG.Cb]).astype(np.uint8)

  # YCrCb -> RGB
  RGB = cv2.cvtColor(YCrCb2, cv2.COLOR_YCrCb2RGB).astype(int)

  return RGB, JPEG

if __name__ == "__main__":
  files = ["0001.jpg", "0002.jpg", "0003.jpg"]
  fragmentsX = [0, 700, 2500]
  fragmentsY = [0, 1900, 3000]

  document = Document()
  document.add_heading('Systemu multimedialne - Lab 7', 1)

  for i in range(3):
    document.add_heading("Zdjecie: {}".format(files[i]), 2)

    for ii in range(2):
      for iii in range(2):
        for iiii in range(3):
          img_orig = cv2.imread(files[i])[fragmentsX[iiii]:fragmentsX[iiii]+128, fragmentsY[iiii]:fragmentsY[iiii]+128]
          img_orig = cv2.cvtColor(img_orig, cv2.COLOR_BGR2RGB)

          fig, axs = plt.subplots(4, 2)
          fig.set_size_inches(9, 13)
          fig.suptitle("{} fragment {} / {} / QY={}, QC={}".format(files[i], iiii, "4:2:2" if ii == 0 else "4:4:4", "QY" if iii == 0 else "1s", "QC" if iii == 0 else "1s"), fontsize=16)
          axs[0, 0].set_title("Original")
          axs[0, 1].set_title("JPEG compression")
          img_YCrCb = cv2.cvtColor(img_orig, cv2.COLOR_RGB2YCrCb).astype(int)
          Y_orig, Cr_orig, Cb_orig = np.clip(cv2.split(img_YCrCb), 0, 255)
          axs[0, 0].imshow(img_orig)
          axs[1, 0].imshow(Y_orig, cmap=plt.cm.gray)
          axs[2, 0].imshow(Cr_orig, cmap=plt.cm.gray)
          axs[3, 0].imshow(Cb_orig, cmap=plt.cm.gray)

          if ii == 0:
            if iii == 0:
              data = CompressJPEG(img_orig, "4:4:4", QY, QC)
            else:
              data = CompressJPEG(img_orig, "4:4:4")
          else:
            if iii == 0:
              data = CompressJPEG(img_orig, "4:2:2", QY, QC)
            else:
              data = CompressJPEG(img_orig, "4:2:2")


          img_decomp, data = DecompressJPEG(data)

          axs[0, 1].imshow(img_decomp)
          axs[1, 1].imshow(data.Y, cmap=plt.cm.gray)
          axs[2, 1].imshow(data.Cr, cmap=plt.cm.gray)
          axs[3, 1].imshow(data.Cb, cmap=plt.cm.gray)

          memfile = BytesIO()
          plt.savefig(memfile)

          document.add_heading("Zdjecie: {} // Współrzędne fragmentu: {}x{} // Chrominacja: {} // QY: {} // QC: {}".format(files[i], fragmentsX[iiii], fragmentsY[iiii], "4:2:2" if ii == 0 else "4:4:4", "QY" if iii == 0 else "1s", "QC" if iii == 0 else "1s"), 3)
          document.add_picture(memfile, width=Inches(5.5))
          memfile.close()
          document.add_paragraph("Oryginał: Warstwa Y = {} / Warstwa Cr = {} / Warstwa Cb = {}\nKompresja: Warstwa Y = {} / Warstwa Cr = {} / Warstwa Cb = {}".format(get_human_size(get_size(Y_orig)), get_human_size(get_size(Cr_orig)), get_human_size(get_size(Cb_orig)), get_human_size(get_size(data.Y[1])), get_human_size(get_size(data.Cr[1])), get_human_size(get_size(data.Cb[1]))))

  document.save('lab7.docx')