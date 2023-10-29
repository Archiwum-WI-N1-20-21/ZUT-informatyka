import matplotlib.pyplot as plt
import numpy as np
import soundfile as sf
import scipy.fftpack
from docx import Document
from docx.shared import Inches
from io import BytesIO

def plotAudio(Signal,Fs,fsize, TimeMargin=[0,0.02]):
  fig ,axs = plt.subplots(2,1,figsize=(10,7))
  fig.tight_layout(pad=3)
  plt.subplot(2,1,1)
  plt.plot(np.arange(0,data.shape[0])/fs,data)
  plt.xlim(TimeMargin[0], TimeMargin[1])
  plt.title("Sygnał dźwiękowy")
  plt.xlabel("Czas (s)")
  plt.subplot(2,1,2)
  yf = scipy.fftpack.fft(data,fsize)
  plt.plot(np.arange(0,fs/2,fs/fsize),20*np.log10(np.abs(yf[:fsize//2])))
  plt.title("Widmo")
  plt.xlabel("Częstotliwość (Hz)")
  plt.ylabel("Natężenie (dB)")

  try:
    maxVal = np.argmax(20*np.log10(np.abs(yf[:fsize//2])))
  except ZeroDivisionError:
    maxVal = 0

  return fig, maxVal

if __name__ == '__main__':
  document = Document()
  document.add_heading('Systemu multimedialne - Lab 1 - Zad 3', 0) # tworzenie nagłówków druga wartość to poziom nagłówka

  files = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav']
  fsizes = [2**8, 2**12, 2**16]

  for file in files:
      document.add_heading('Plik - {}'.format(file),2)
      for i,fsize in enumerate(fsizes):
          document.add_heading('fsize - {}'.format(fsize ),3) # nagłówek sekcji, mozę być poziom wyżej

          ############################################################
          # Tu wykonujesz jakieś funkcje i rysujesz wykresy
          ############################################################

          data, fs = sf.read(files[i], dtype=np.int32)

          memfile = BytesIO()
          f, maxVal = plotAudio(data, fs, fsize)
          f.savefig(memfile)
          document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku

          memfile.close()
          ############################################################
          # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
          document.add_paragraph('Wartość maksymalna = {}'.format(maxVal))
          ############################################################

  document.save('report.docx') # zapis do pliku