import pandas as pd
import cv2
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from io import BytesIO

def convertImage(img):
  Y1 = np.dot(img[...,:3], [0.299, 0.587, 0.114])
  Y2 = np.dot(img[...,:3], [0.2126, 0.7152, 0.0722])

  # Czerwony
  img_red_gray = img.copy()
  img_red_gray = img[..., 0]

  img_red = img.copy()
  img_red[..., 1] = 0
  img_red[..., 2] = 0

  # Zielony
  img_green_gray = img.copy()
  img_green_gray = img[..., 1]

  img_green = img.copy()
  img_green[..., 0] = 0
  img_green[..., 2] = 0

  # Niebieski
  img_blue_gray = img.copy()
  img_blue_gray = img[..., 2]

  img_blue = img.copy()
  img_blue[..., 0] = 0
  img_blue[..., 1] = 0

  fig ,axs = plt.subplots(3,3)
  fig.tight_layout(pad=3)
  plt.subplot(3,3,1)
  plt.title("Oryginalny")
  plt.imshow(img)
  plt.subplot(3,3,2)
  plt.title("Y1")
  plt.imshow(Y1, cmap='gray')
  plt.subplot(3,3,3)
  plt.title("Y2")
  plt.imshow(Y2, cmap='gray')
  plt.subplot(3,3,4)
  plt.title("Gray R")
  plt.imshow(img_red_gray, cmap='gray')
  plt.subplot(3,3,5)
  plt.title("Gray G")
  plt.imshow(img_green_gray, cmap='gray')
  plt.subplot(3,3,6)
  plt.title("Gray B")
  plt.imshow(img_blue_gray, cmap='gray')
  plt.subplot(3,3,7)
  plt.title("Color R")
  plt.imshow(img_red)
  plt.subplot(3,3,8)
  plt.title("Color G")
  plt.imshow(img_green)
  plt.subplot(3,3,9)
  plt.title("Color B")
  plt.imshow(img_blue)

  return fig

if __name__ == '__main__':
  document = Document()
  document.add_heading('Systemu multimedialne - Lab 2 - Zad 3', 0)

  df = pd.DataFrame()

  df = pd.DataFrame(data={'Filename':['B02.jpg'],
                          'Fragments':[[[0,200,0,200],[200,400,200,400],[400,600,400,600]]]
                          })

  print(df)

  for index, row in df.iterrows():
      img = plt.imread(row['Filename'])
      document.add_heading('Plik - {}'.format(row['Filename']),2)
      if row['Fragments'] is not None:
          # mamy nie pustą listę fragmentów
          for f in row['Fragments']:
              document.add_heading('Współrzędne - {}'.format(f),3)

              fragment = img[f[0]:f[1],f[2]:f[3]].copy()

              memfile = BytesIO()
              f = convertImage(fragment)
              f.savefig(memfile)

              document.add_picture(memfile, width=Inches(6))

              memfile.close()

  document.save('report.docx')
