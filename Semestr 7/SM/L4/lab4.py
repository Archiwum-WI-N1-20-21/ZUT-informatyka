import numpy as np
import soundfile as sf
import sounddevice as sd
import matplotlib.pyplot as plt
import scipy.fftpack
from scipy.interpolate import interp1d
from docx import Document
from docx.shared import Inches
from io import BytesIO


def plotAudio(data, fs, TimeMargin=np.array([0, 0.02]), Name="", fsize=2**8):
  fig, (ax1, ax2) = plt.subplots(2, constrained_layout=True, figsize=(10, 7))
  fig.suptitle(Name)
  ax1.set_xlim(TimeMargin)
  ax1.set_xlabel("Czas [s]")
  ax1.set_title("Sygnał dźwiękowy")
  ax1.plot(np.arange(0, data.shape[0]) / fs, data)
  ax2.set_title("Widmo")
  ax2.set_xlabel("Częstotliwość [Hz]")
  ax2.set_ylabel("Natężenie [dB]")
  yf = scipy.fftpack.fft(data, fsize)
  yf = yf + np.finfo(np.float32).eps
  ax2.plot(np.arange(0, fs / 2, fs / fsize), 20 * np.log10(np.abs(yf[: fsize // 2])))
  return fig


def Kwant(data, bit):
  dataF = data.copy()

  d = (2 ** bit) - 1

  if np.issubdtype(data.dtype, np.floating):
    m = -1
    n = 1
  else:
    dataF =dataF.astype(np.float32)
    m = np.iinfo(data_input.dtype).min
    n = np.iinfo(data_input.dtype).max

  dataF = dataF - m
  dataF = dataF / (n-m)
  dataF = dataF * d
  dataF = np.round(dataF)
  dataF = dataF / d
  dataF = dataF * (n-m)
  dataF = dataF + m

  return dataF.astype(data.dtype)


def Decimation(data, n):
  return data[::n]

def Interpolation(data, new_length, kind='linear'):
    old_indices = np.arange(len(data))
    new_indices = np.linspace(0, len(data) - 1, new_length)
    interpolator = interp1d(old_indices, data, kind=kind)
    return interpolator(new_indices)

if __name__ == "__main__":
  document = Document()
  document.add_heading('Systemu multimedialne - Lab 4', 0)

  sinFiles = ["SIN/sin_60Hz.wav", "SIN/sin_440Hz.wav", "SIN/sin_8000Hz.wav", "SIN/sin_combined.wav"]
  bits = [4, 8, 16, 24]
  decSteps = [2, 4, 6, 10, 24]
  frequencies = [2000, 4000, 8000, 11999, 16000, 16953, 24000, 41000]
  linear = [True, False]

  '''for sinFile in sinFiles:
    document.add_heading("Plik: " + sinFile, 2)
    data, fs = sf.read(sinFile, dtype=np.single)

    for bit in bits:
      print("Processing: {} // Bit: {}".format(sinFile, bit))
      document.add_heading("Plik: {}, Bit: {}".format(sinFile, bit), 3)
      kwantData = Kwant(data, bit)
      memfile = BytesIO()
      fig = plotAudio(kwantData, fs)
      fig.savefig(memfile)
      document.add_picture(memfile, width=Inches(6))
      memfile.close()
      plt.close()

    for step in decSteps:
      print("Processing: {} // Step: {}".format(sinFile, step))
      document.add_heading("Plik: {}, Krok: {}".format(sinFile, step), 3)
      decData = Decimation(data, step)
      memfile = BytesIO()
      fig = plotAudio(decData, fs)
      fig.savefig(memfile)
      document.add_picture(memfile, width=Inches(6))
      memfile.close()
      plt.close()

    for f in frequencies:
      for isLinear in linear:
        print("Processing: {} // Frequency: {} // Linear: {}".format(sinFile, f, isLinear))
        if isLinear:
          document.add_heading("Plik: {}, Częstotliwość: {}, Interpolacja: {}".format(sinFile, f, "Liniowa"), 3)
          interpData = Interpolation(data, f, "linear")
        else:
          document.add_heading("Plik: {}, Częstotliwość: {}, Interpolacja: {}".format(sinFile, f, "Nieliniowa"), 3)
          interpData = Interpolation(data, f, "nearest")

        memfile = BytesIO()
        fig = plotAudio(interpData, fs)
        fig.savefig(memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        plt.close()

  document.save('lab4.docx')'''

  singFiles = ["SING/sing_high1.wav", "SING/sing_low2.wav", "SING/sing_medium1.wav"]
  bits = [4, 8]
  decSteps = [4, 6, 10, 24]
  frequencies = [4000, 8000, 11999, 16000, 16953]
  linear = [True, False]

  for singFile in singFiles:
    data, fs = sf.read(singFile, dtype=np.single)

    for bit in bits:
      print("Processing: {} // Bit: {}".format(singFile, bit))
      data = Kwant(data, bit)
      sd.play(data, fs)
      sf.write("{}_{}.wav".format(singFile, bit), data, fs)

    for step in decSteps:
      print("Processing: {} // Step: {}".format(singFile, step))
      decData = Decimation(data, step)
      sd.play(decData, fs)
      sf.write("{}_{}.wav".format(singFile, step), decData, fs)

    for f in frequencies:
      for isLinear in linear:
        print("Processing: {} // Frequency: {} // Linear: {}".format(singFile, f, isLinear))

        if isLinear:
          interpData = Interpolation(data, f, "linear")
        else:
          interpData = Interpolation(data, f, "nearest")
        sd.play(interpData, f)

        sf.write("{}_{}_{}.wav".format(singFile, f, isLinear), interpData, f)
