import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from io import BytesIO

np.set_printoptions(threshold=np.inf)

pallet1 = np.linspace(0,1,2).reshape(2,1)
pallet2 = np.linspace(0,1,4).reshape(4,1)
pallet4 = np.linspace(0,1,16).reshape(16,1)

pallet8 = np.array([
        [0.0, 0.0, 0.0,],
        [0.0, 0.0, 1.0,],
        [0.0, 1.0, 0.0,],
        [0.0, 1.0, 1.0,],
        [1.0, 0.0, 0.0,],
        [1.0, 0.0, 1.0,],
        [1.0, 1.0, 0.0,],
        [1.0, 1.0, 1.0,]])

pallet16 =  np.array([
        [0.0, 0.0, 0.0,],
        [0.0, 1.0, 1.0,],
        [0.0, 0.0, 1.0,],
        [1.0, 0.0, 1.0,],
        [0.0, 0.5, 0.0,],
        [0.5, 0.5, 0.5,],
        [0.0, 1.0, 0.0,],
        [0.5, 0.0, 0.0,],
        [0.0, 0.0, 0.5,],
        [0.5, 0.5, 0.0,],
        [0.5, 0.0, 0.5,],
        [1.0, 0.0, 0.0,],
        [0.75, 0.75, 0.75,],
        [0.0, 0.5, 0.5,],
        [1.0, 1.0, 1.0,],
        [1.0, 1.0, 0.0,]])

def colorFit(color,Pallet):
  return Pallet[np.argmin(np.linalg.norm(Pallet-color,axis=1))]

def kwant_colorFit(img,Pallet):
  out_img = img.copy()

  for w in range(out_img.shape[0]):
    for k in range(out_img.shape[1]):
      out_img[w,k] = colorFit(img[w,k],Pallet)

  return out_img

def random_dithering(img):
  out_img = img.copy()

  r = np.random.rand(*img.shape)

  if len(img.shape) == 3:
    for w in range(out_img.shape[0]):
      for k in range(out_img.shape[1]):
        for c in range(out_img.shape[2]):
          if out_img[w][k][c] > r[w][k][c]:
            out_img[w][k][c] = 1
          else:
            out_img[w][k][c] = 0
  elif len(img.shape) == 2:
    for w in range(out_img.shape[0]):
      for k in range(out_img.shape[1]):
        if out_img[w][k] > r[w][k]:
          out_img[w][k] = 1
        else:
          out_img[w][k] = 0

  return out_img

def create_bayer_matrix(n = 2):
  if n == 1:
    return np.array([[0, 2], [3, 1]])
  else:
    smaller_matrix = create_bayer_matrix(n - 1)
    return np.block([[4 * smaller_matrix, 4 * smaller_matrix + 2],
                     [4 * smaller_matrix + 3, 4 * smaller_matrix + 1]])

def create_mpre(matrix):
    n = np.log2(matrix.shape[0])
    return (matrix + 1) / (2 * n)**2 - 0.5

def ordered_dithering(img, pallete, N = 2):
  out_img = img.copy()

  bayer_matrix = create_bayer_matrix(N)
  mpre = create_mpre(bayer_matrix)

  N = 2 ** N
  for i in range(img.shape[0]):
    for j in range(img.shape[1]):
      tmpPixelVal = img[i, j] + mpre[i % N, j % N]
      out_img[i, j] = colorFit(tmpPixelVal, pallete)

  return out_img

def floyd_steinberg_dithering(img, pallete):
  out_img = img.copy()

  height = out_img.shape[0]
  width = out_img.shape[1]

  for y in range(height):
    for x in range(width):
      oldpixel = out_img[y, x]
      newpixel = colorFit(oldpixel, pallete)
      quant_error = oldpixel - newpixel
      out_img[y, x] = newpixel

      if x + 1 < width:
        out_img[y][x + 1] += quant_error * 7 / 16
      if y + 1 < height and x - 1 >= 0:
        out_img[y + 1][x - 1] += quant_error * 3 / 16
      if y + 1 < height:
        out_img[y + 1][x] += quant_error * 5 / 16
      if y + 1 < height and x + 1 < width:
        out_img[y + 1][x + 1] += quant_error * 1 / 16

  return out_img

def make_plot(img, is1bit, pallet, isGray):
  if img.dtype == "uint8":
    img = img / 255.0

  if is1bit:
    fig ,axs = plt.subplots(1,5, figsize=(24, 10))
  else:
    fig ,axs = plt.subplots(1,4, figsize=(24, 10))

  plt.tight_layout(pad=0.4, w_pad=0.5, h_pad=1.0)

  # Oryginalny
  if is1bit:
    plt.subplot(1,5,1)
  else:
    plt.subplot(1,4,1)
  plt.title("Oryginalny", fontsize=18)
  if isGray:
    plt.imshow(img, cmap='gray')
  else:
    plt.imshow(img)

  # Kwantyzacja
  if is1bit:
    plt.subplot(1,5,2)
  else:
    plt.subplot(1,4,2)
  plt.title("Kwantyzacja", fontsize=18)
  img_kwantyzacja = kwant_colorFit(img,pallet)
  if isGray:
    plt.imshow(img_kwantyzacja, cmap='gray')
  else:
    plt.imshow(img_kwantyzacja)

  # Dithering losowy
  if is1bit:
    plt.subplot(1,5,3)
    plt.title("Dithering\nlosowy", fontsize=18)
    img_dith_losowy = random_dithering(img)
    plt.imshow(img_dith_losowy, cmap='gray')

  # Dithering zorganizowany
  if is1bit:
    plt.subplot(1,5,4)
  else:
    plt.subplot(1,4,3)
  plt.title("Dithering\nzorganizowany", fontsize=18)
  img_dith_zorganizowany = ordered_dithering(img, pallet)
  if isGray:
    plt.imshow(img_dith_zorganizowany, cmap='gray')
  else:
    plt.imshow(img_dith_zorganizowany)

  # Dithering Floyd-Steinberg
  if is1bit:
    plt.subplot(1,5,5)
  else:
    plt.subplot(1,4,4)
  plt.title("Dithering\nFloyd-Steinberg", fontsize=18)
  img_dith_floyd = floyd_steinberg_dithering(img, pallet)
  if isGray:
    plt.imshow(img_dith_floyd, cmap='gray')
  else:
    plt.imshow(img_dith_floyd)

  return fig

if __name__ == "__main__":
  document = Document()
  document.add_heading('Systemu multimedialne - Lab 3', 0)

  images = ["IMG_GS/GS_0001.tif", "IMG_GS/GS_0002.png", "IMG_GS/GS_0003.png",
            "IMG_SMALL/SMALL_0004.jpg", "IMG_SMALL/SMALL_0006.jpg",
            "IMG_SMALL/SMALL_0007.jpg", "IMG_SMALL/SMALL_0009.jpg"]

  pallets = [pallet1, pallet2, pallet4, pallet8, pallet16]

  for idxImage, image in enumerate(images):
    img = plt.imread(image)
    for idxPallete, pallet in enumerate(pallets):
      if idxImage < 3 and idxPallete > 2:
        continue

      if idxImage > 2 and idxPallete < 3:
        continue

      if idxImage < 3:
        isGray = True

        if (len(img.shape) > 2):
          img = 0.299 * img[:,:,0] + 0.587 * img[:,:,1] + 0.114 * img[:,:,2]
      else:
        isGray = False

        if (img.shape[2]>3):
          img_no_transparency = img[:,:,0:3].copy()
          img=img_no_transparency

      print("Processing: {} // Pallet: {}bit".format(image, 2 ** idxPallete))
      document.add_heading("Zdjecie: {} // Paleta: {}bit".format(image, 2 ** idxPallete))

      memfile = BytesIO()

      if idxPallete == 0:
        f = make_plot(img, True, pallet, isGray)
      else:
        f = make_plot(img, False, pallet, isGray)

      f.savefig(memfile)
      document.add_picture(memfile, width=Inches(6))
      memfile.close()

  document.save('lab3.docx')